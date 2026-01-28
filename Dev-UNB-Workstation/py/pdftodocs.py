# # import pdfplumber
# # from docx import Document

# # def pdf_to_docx(pdf_path, docx_path):
# #     try:
# #         # PDF file ko read karna
# #         with pdfplumber.open(pdf_path) as pdf:
# #             document = Document()
            
# #             # Har page ka text extract karna
# #             for page in pdf.pages:
# #                 text = page.extract_text()
# #                 if text:
# #                     # Text ko Word document mein add karna
# #                     document.add_paragraph(text)
# #                     document.add_paragraph("\n")  # Page break ke liye
                
# #             # DOCX file save karna
# #             document.save(docx_path)
# #             print(f"PDF successfully converted to DOCX: {docx_path}")
# #     except Exception as e:
# #         print(f"Error: {e}")

# # # Input aur output file ka path
# # pdf_path = r"C:\Users\Latitude\Downloads\william wali episode 5.pdf" # PDF ka path
# # docx_path = "C:\\Users\\Latitude\\Downloads\\urdu.docx"  # DOCX file ka output path

# # pdf_to_docx(pdf_path, docx_path)





# import pytesseract
# from pdf2image import convert_from_path
# from docx import Document

# def pdf_to_docx_using_ocr(pdf_path, docx_path):
#     try:
#         # PDF ko images mein convert karna
#         images = convert_from_path(pdf_path)
#         document = Document()
        
#         for image in images:
#             # OCR se text extract karna (Urdu language ke liye)
#             text = pytesseract.image_to_string(image, lang='urd')
#             document.add_paragraph(text)
#             document.add_paragraph("\n")  # Page break ke liye

#         # DOCX file save karna
#         document.save(docx_path)
#         print(f"PDF successfully converted to DOCX: {docx_path}")
#     except Exception as e:
#         print(f"Error: {e}")

# # Input aur output file ka path
# pdf_path = r"C:\Users\Latitude\Downloads\william wali episode 5.pdf" # PDF ka path
# docx_path = "C:\\Users\\Latitude\\Downloads\\urdu.docx"  # DOCX file ka output path

# pdf_to_docx_using_ocr(pdf_path, docx_path)













from pdf2image import convert_from_path
import pytesseract
from docx import Document
import os

# Poppler ka path specify karein
poppler_path = r"C:\Release-24.08.0-0\poppler-24.08.0\Library\bin"
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def pdf_to_docx(pdf_path, output_docx_path):
    # PDF ko images mein convert karna
    try:
        images = convert_from_path(pdf_path, poppler_path=poppler_path)
    except Exception as e:
        print(f"Error: {e}")
        return

    # Word document create karna
    document = Document()

    # Har image ko process karna aur text extract karna
    for i, image in enumerate(images):
        print(f"Processing page {i + 1}/{len(images)}...")
        text = pytesseract.image_to_string(image, lang="urd")
        document.add_paragraph(text)
        document.add_page_break()

    # DOCX save karna
    document.save(output_docx_path)
    print(f"PDF converted to DOCX successfully: {output_docx_path}")

# Input aur output files ka path
pdf_path = r"C:\Users\Latitude\Downloads\william wali episode 5.pdf"
output_docx_path = "C:\Users\Latitude\Downloads\william wali episode 5 doxxxx.pdf"

# Function call
pdf_to_docx(pdf_path, output_docx_path)
